"""Document → text parsing (Phase 4 Step 2). Pure, no LLM."""
import io

import pytest

from integrations.document_parser import (
    DocumentParseError, MAX_BYTES, MAX_CHARS, parse_text, parse_upload,
)


def test_parse_text_ok_and_too_short():
    assert parse_text("  Week 1: read chapter one and take notes  ").startswith("Week 1")
    with pytest.raises(DocumentParseError) as e:
        parse_text("hi")
    assert e.value.code == "too_short"


def test_parse_text_truncates_when_too_long():
    huge = "a" * (MAX_CHARS + 500)
    assert len(parse_text(huge)) == MAX_CHARS


def test_parse_upload_txt_and_md():
    body = "Week 1: read chapter one\nWeek 2: problem set".encode("utf-8")
    assert "Week 2" in parse_upload("syllabus.txt", body)
    assert "Week 2" in parse_upload("syllabus.md", body)


def test_parse_upload_docx_paragraphs_and_tables():
    from docx import Document

    doc = Document()
    doc.add_paragraph("Week 1: read chapter one")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Task"
    table.rows[0].cells[1].text = "Due"
    buf = io.BytesIO()
    doc.save(buf)

    text = parse_upload("syllabus.docx", buf.getvalue())
    assert "Week 1: read chapter one" in text
    assert "Task | Due" in text


def test_parse_upload_unsupported_type():
    with pytest.raises(DocumentParseError) as e:
        parse_upload("photo.gif", b"GIF89a" + b"x" * 40)
    assert e.value.code == "unsupported"


def test_parse_upload_too_large():
    with pytest.raises(DocumentParseError) as e:
        parse_upload("big.txt", b"x" * (MAX_BYTES + 1))
    assert e.value.code == "too_large"


def test_parse_upload_bad_pdf():
    # Sniffed as PDF by the %PDF magic but unreadable → structured error, no crash.
    with pytest.raises(DocumentParseError) as e:
        parse_upload("broken.pdf", b"%PDF-1.4 not really a pdf")
    assert e.value.code in ("pdf_unreadable", "pdf_no_text")
