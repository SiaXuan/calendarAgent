"""
Document → plain text (Phase 4 Step 2).

Cheap, deterministic pre-LLM stage: sniff the format, pull text, enforce hard
size/length limits, and reject empty / too-short input before spending a Claude
call. Supported this pass: pasted text, .md/.txt, .pdf (pypdf), .docx
(python-docx). Images/vision are deferred (would add Pillow + a vision call).
"""
import io
import logging

# pypdf logs a "Ignoring wrong pointing object …" warning for every malformed
# cross-reference pointer it recovers from. Harmless (it still extracts the text)
# but noisy — quiet it to errors only.
logging.getLogger("pypdf").setLevel(logging.ERROR)

MAX_BYTES = 10 * 1024 * 1024      # 10 MB upload ceiling
MAX_PDF_PAGES = 30
MAX_CHARS = 40_000                # ~ what fits comfortably in one extraction call
MIN_CHARS = 20                    # below this there's nothing to plan from


class DocumentParseError(Exception):
    """Recoverable parse/validation failure → the API maps it to HTTP 422."""

    def __init__(self, code: str, message: str):
        self.code = code
        self.message = message
        super().__init__(message)


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001 — malformed upload, not our bug
        raise DocumentParseError("pdf_unreadable", f"Could not read PDF: {e}")
    if len(reader.pages) > MAX_PDF_PAGES:
        raise DocumentParseError(
            "pdf_too_long",
            f"PDF has {len(reader.pages)} pages (max {MAX_PDF_PAGES}).",
        )
    parts: list[str] = []
    for page in reader.pages:
        # "layout" mode keeps columns/tables roughly aligned (course schedules,
        # PRD tables) — the default mode interleaves them into unreadable text
        # the LLM then rejects as "not a plan". Fall back per-page on any error
        # so one bad page can't crash the whole import.
        try:
            parts.append(page.extract_text(extraction_mode="layout") or "")
        except Exception:  # noqa: BLE001
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue
    text = "\n".join(parts)
    if not text.strip():
        # A scanned/image-only PDF yields no text — no OCR this pass.
        raise DocumentParseError(
            "pdf_no_text",
            "No text found in the PDF (looks scanned). Paste the text or a screenshot instead.",
        )
    return text


def _extract_docx(data: bytes) -> str:
    from docx import Document

    try:
        doc = Document(io.BytesIO(data))
    except Exception as e:  # noqa: BLE001
        raise DocumentParseError("docx_unreadable", f"Could not read DOCX: {e}")
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _validate_text(text: str) -> str:
    text = text.strip()
    if len(text) < MIN_CHARS:
        raise DocumentParseError(
            "too_short", "The document is empty or too short to plan from.")
    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS]     # truncate rather than reject; keep the head
    return text


def parse_text(text: str) -> str:
    """Pasted-text path."""
    return _validate_text(text)


def parse_upload(filename: str, data: bytes) -> str:
    """
    Uploaded-file path. Routes by extension (magic-sniffed as a fallback) and
    enforces the byte ceiling. Raises DocumentParseError on anything unusable.
    """
    if len(data) > MAX_BYTES:
        raise DocumentParseError(
            "too_large", f"File is larger than {MAX_BYTES // (1024 * 1024)} MB.")

    name = (filename or "").lower()
    if name.endswith((".txt", ".md")):
        try:
            raw = data.decode("utf-8")
        except UnicodeDecodeError:
            raise DocumentParseError("bad_encoding", "Text file is not valid UTF-8.")
    elif name.endswith(".pdf") or data[:4] == b"%PDF":
        raw = _extract_pdf(data)
    elif name.endswith(".docx") or data[:4] == b"PK\x03\x04":
        raw = _extract_docx(data)
    else:
        raise DocumentParseError(
            "unsupported",
            "Unsupported file type. Use .txt, .md, .pdf, or .docx (or paste the text).",
        )
    return _validate_text(raw)
